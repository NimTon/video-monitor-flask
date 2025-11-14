import cv2
import time
import numpy as np
from datetime import datetime
from PIL import Image
import io
from pypinyin import lazy_pinyin
from docx import Document, oxml
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
import os
import base64
import asyncio
from typing import List, Tuple
import pythoncom
from win32com import client
import re
import pandas as pd
from utils.log_utils import log


def get_first_changed_row(df):
    """
    按 timestamp 每 10 秒间隔取第一行。
    要求 df 中有列 'timestamp'（datetime 类型）。
    """
    if df.empty or 'timestamp' not in df.columns:
        raise ValueError("输入数据缺少 timestamp 列")

    # 确保时间列是 datetime 类型
    df = df.copy()
    df['timestamp'] = pd.to_datetime(df['timestamp'])
    df = df.sort_values('timestamp')

    # 计算分组标签（每 10 秒一个分组）
    df['time_group'] = (df['timestamp'].astype('int64') // 10_000_000_000)  # 每 10 秒一组
    # 取每组第一条
    result = df.groupby('time_group', as_index=False).first()

    # 去掉辅助列
    result = result.drop(columns=['time_group'])
    return result


def camel_to_snake(name: str) -> str:
    """驼峰转下划线"""
    s1 = re.sub('(.)([A-Z][a-z]+)', r'\1_\2', name)
    return re.sub('([a-z0-9])([A-Z])', r'\1_\2', s1).lower()


async def clean_task(paths, interval=24 * 3600, days=7):
    """
    异步定时清理任务
    - 启动时立即执行一次
    - 之后每隔 interval 秒执行一次
    """
    while True:
        log("INFO", f"开始清理旧文件, {', '.join(paths)}")
        for path in paths:
            cleared_path = clean_old_files(path, days)
            if cleared_path:
                log("INFO", f"清理了 {path} 中 {len(cleared_path)} 个过期文件")
        # 首次执行完直接等待 interval 秒，再执行下一轮
        await asyncio.sleep(interval)


def clean_old_files(path: str, day: int):
    """
    清理 path 下修改时间距今超过 day 天的文件（不删除目录）
    :param path: 文件夹路径
    :param day: 天数阈值
    :return: 被删除文件的路径列表
    """
    if not os.path.exists(path):
        return []

    now = time.time()
    cutoff = now - day * 86400  # 秒
    removed_files = []

    for root, dirs, files in os.walk(path):
        log("INFO", f"检查目录 {root}")
        for filename in files:
            file_path = os.path.join(root, filename)
            try:
                log("INFO", f"检查文件 {file_path}")
                if os.path.isfile(file_path) and os.path.getmtime(file_path) < cutoff:
                    log("INFO", f"删除文件 {file_path}")
                    os.remove(file_path)
                    removed_files.append(file_path)
                else:
                    log("INFO", f"文件 {file_path} 不满足条件")
            except Exception:
                log("FAIL", f"清理文件 {file_path} 时出错")

    return removed_files


def to_png_bytes(img: np.ndarray) -> bytes:
    """
    将 OpenCV 图像编码为 PNG 字节流

    参数:
    - img: np.ndarray，BGR 或 BGRA 图像

    返回:
    - bytes，PNG 格式字节流
    """
    success, buf = cv2.imencode(".png", img)
    if not success:
        raise ValueError("PNG 编码失败")
    return buf.tobytes()


def capture_frame_from_url(url: str) -> np.ndarray:
    """
    从视频流 URL 捕获一帧
    """
    cap = cv2.VideoCapture(url)
    if not cap.isOpened():
        raise ValueError(f"无法打开视频流: {url}")
    ret, frame = cap.read()
    cap.release()
    if not ret:
        raise ValueError(f"无法从视频流获取帧: {url}")
    return frame


def relative_to_pixel_fence(
        url: str,
        relative_fence_points: List[dict],
        target_height: int = None
) -> Tuple[np.ndarray, List[Tuple[int, int]]]:
    """
    将归一化 fence_points 转换为像素坐标，可选择缩放到固定高度保持宽高比

    Args:
        url: 视频流 URL
        relative_fence_points: [{"x": x_rel, "y": y_rel}, ...] 归一化坐标
        target_height: 输出帧高度，可选。不传则返回原始帧

    Returns:
        frame: 视频帧 np.ndarray
        pixel_points: [(x_pixel, y_pixel), ...] 对应像素坐标
    """
    frame = capture_frame_from_url(url)
    if frame is None or frame.size == 0:
        raise ValueError(f"无法从 {url} 获取有效帧")

    orig_h, orig_w = frame.shape[:2]

    if target_height is not None:
        # 缩放比例
        scale = target_height / orig_h
        new_w = int(orig_w * scale)
        new_h = target_height
        frame_resized = cv2.resize(frame, (new_w, new_h), interpolation=cv2.INTER_AREA)
        w, h = new_w, new_h
        frame_out = frame_resized
    else:
        w, h = orig_w, orig_h
        frame_out = frame

    # 生成像素坐标
    pixel_points = []
    for p in relative_fence_points:
        try:
            x_rel = float(p.get("x", 0))
            y_rel = float(p.get("y", 0))
        except (ValueError, TypeError):
            x_rel, y_rel = 0, 0

        x_pixel = max(0, min(int(x_rel * w), w - 1))
        y_pixel = max(0, min(int(y_rel * h), h - 1))
        pixel_points.append((x_pixel, y_pixel))

    return frame_out, pixel_points


def docx_to_pdf(docx_path: str, pdf_path: str = None) -> str:
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"{docx_path} 不存在")

    if pdf_path is None:
        pdf_path = os.path.splitext(docx_path)[0] + ".pdf"

    try:
        # 初始化 COM
        pythoncom.CoInitialize()

        word = client.Dispatch('Word.Application')  # 或 CreateObject 也行
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 表示 PDF 格式
        doc.Close()
        word.Quit()
        return pdf_path
    except Exception as e:
        raise RuntimeError(f"导出 PDF 失败: {e}")
    finally:
        # 释放 COM
        pythoncom.CoUninitialize()


def save_report_to_docx(content: str, save_dir: str, filename: str, title: str = None, images: list = None,
                        image_captions: list = None):
    """
    保存报告到 Word 文件（docx），带中文/英文字体设置、标题、正文、图片及描述。
    :param content: 正文内容
    :param save_dir: 保存目录
    :param filename: 文件名
    :param title: 标题文本
    :param images: 图片路径列表
    :param image_captions: 图片描述列表，长度应与 images 一致
    :return: 文件保存路径
    """
    try:
        os.makedirs(save_dir, exist_ok=True)
        doc = Document()

        # ===== 标题 =====
        if title:
            heading = doc.add_heading(level=1)
            run = heading.add_run(title)
            run.font.size = Pt(22)  # 小二
            run.bold = True
            run.font.name = 'Times New Roman'  # 英文
            run._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), '宋体')  # 中文
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            # 设置行距1.5倍
            heading.paragraph_format.line_spacing = 1.5
            heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中

        # ===== 正文 =====
        para = doc.add_paragraph()
        run = para.add_run(content)
        run.font.size = Pt(12)  # 小四
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), '宋体')
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        para.paragraph_format.line_spacing = 1.0  # 单倍行距
        para.paragraph_format.first_line_indent = Cm(0.74)  # 约2字符缩进
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # ===== 插入图片 =====
        if images:
            for idx, img_path in enumerate(images):
                if os.path.exists(img_path):
                    try:
                        pic = doc.add_picture(img_path, width=Cm(12))  # 宽度固定12cm
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        # 图片描述
                        if image_captions and idx < len(image_captions):
                            caption_para = doc.add_paragraph()
                            caption_run = caption_para.add_run(image_captions[idx])
                            caption_run.font.size = Pt(9)  # 小五
                            caption_run.font.name = 'Times New Roman'
                            caption_run._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), '宋体')
                            caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    except Exception as e:
                        log("FAIL", f"插入图片失败 {img_path}: {e}")

        # 保存文件
        filepath = os.path.join(save_dir, filename)
        doc.save(filepath)
        log("SUCCESS", f"Word 报告已保存: {filepath}")
        return filepath
    except Exception as e:
        log("FAIL", f"保存 Word 报告失败: {filename}, 错误: {e}")
        return None


def points_to_abs_points(frame, fences):
    height, width = frame.shape[:2]
    fence_points = []
    for fence in fences:
        points = fence.get('points', [])
        abs_points = [(int(p['x'] * width), int(p['y'] * height)) for p in points]
        fence_points.append(abs_points)
    return fence_points


# 在报警分发时，给frame加上红色围栏标记
def draw_fence_on_frame(frame, fence_points, changed=None):
    """
    在 frame 上绘制围栏，根据 changed 的 bool 自动切换颜色：
        True  -> 红色
        False -> 绿色
    兼容 BGR / BGRA 图像
    """
    if not fence_points or len(fence_points) < 3:
        return frame

    # 根据 changed 设置颜色
    # BGR 或 BGRA（若是 BGRA 会自动补上 alpha）
    if changed is True:
        color_bgr = (0, 0, 255)      # red
        color_bgra = (0, 0, 255, 255)
    elif changed is False:
        color_bgr = (0, 255, 0)      # green
        color_bgra = (0, 255, 0, 255)
    else:
        color_bgr = (0, 0, 255)
        color_bgra = (0, 0, 255, 255)

    pts = np.array(fence_points, np.int32).reshape((-1, 1, 2))

    # -------- BGR 图像 --------
    if frame.shape[2] == 3:
        cv2.polylines(frame, [pts], isClosed=True, color=color_bgr, thickness=2)
        for (x, y) in fence_points:
            cv2.circle(frame, (x, y), radius=4, color=color_bgr, thickness=-1)

    # -------- BGRA 图像（带透明度）--------
    elif frame.shape[2] == 4:
        overlay = np.zeros_like(frame, dtype=np.uint8)

        # 在 overlay 层绘制
        cv2.polylines(overlay, [pts], isClosed=True, color=color_bgra, thickness=2)
        for (x, y) in fence_points:
            cv2.circle(overlay, (x, y), radius=4, color=color_bgra, thickness=-1)

        # 覆盖非透明部分
        mask = overlay[:, :, 3] > 0
        frame[mask] = overlay[mask]

    return frame


def cv2_frame_to_base64(frame):
    # 转成 RGB
    rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
    pil_img = Image.fromarray(rgb_frame)
    buffered = io.BytesIO()
    pil_img.save(buffered, format="JPEG")
    img_str = base64.b64encode(buffered.getvalue()).decode()
    return img_str


def image_path_to_base64(image_path: str) -> str:
    if not os.path.exists(image_path):
        raise FileNotFoundError(f"文件不存在: {image_path}")
    with open(image_path, "rb") as f:
        img_bytes = f.read()
    return base64.b64encode(img_bytes).decode("utf-8")


# MD5哈希函数
def md5(str):
    import hashlib
    m = hashlib.md5()  # 创建MD5哈希对象
    m.update(str.encode("utf8"))  # 更新哈希内容
    return m.hexdigest()  # 返回16进制哈希值

def save_imgs_as_video(frame_paths, output_path, fps=25):
    """
    将一系列图像路径合并为一个视频文件
    
    Args:
        frame_paths: 图像文件路径列表
        output_path: 输出视频文件路径
        fps: 视频帧率，默认25
    
    Returns:
        str: 输出视频文件路径
    """
    if not frame_paths:
        raise ValueError("没有图像路径提供")
    
    # 读取第一帧以获取尺寸
    first_frame = cv2.imread(frame_paths[0])
    if first_frame is None:
        raise ValueError(f"无法读取第一帧图像: {frame_paths[0]}")
    
    height, width = first_frame.shape[:2]
    fourcc = cv2.VideoWriter_fourcc(*'H264')
    
    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # 创建VideoWriter对象
    video_writer = cv2.VideoWriter(output_path, fourcc, fps, (width, height))
    
    # 写入每一帧
    for frame_path in frame_paths:
        frame = cv2.imread(frame_path)
        if frame is None:
            print(f"警告: 无法读取帧 {frame_path}，跳过")
            continue
            
        # 如果尺寸不匹配，调整尺寸
        if frame.shape[1] != width or frame.shape[0] != height:
            frame = cv2.resize(frame, (width, height))
            
        video_writer.write(frame)
    
    video_writer.release()
    return output_path

def save_frames_as_video(stream_id, fence_id, frames, video_root='./videos', base_url='127.0.0.1:5000', fps=25):
    """
    直接在 video_root 下生成 MP4 文件，不创建子文件夹。
    文件名格式: {stream_id}_{fence_id}_{时间戳}.mp4
    """
    now_time_str = datetime.now().strftime('%Y%m%d_%H%M%S')
    os.makedirs(f"{video_root}/{stream_id}", exist_ok=True)

    # 删除超过7天的旧文件
    # for file in os.listdir(video_root):
    #     file_path = os.path.join(video_root, file)
    #     if os.path.isfile(file_path):
    #         try:
    #             # 从文件名解析时间戳
    #             ts = file.split('_')[-1].split('.')[0]
    #             file_time = datetime.strptime(ts, '%Y%m%d_%H%M%S')
    #             if now - file_time > timedelta(days=7):
    #                 os.remove(file_path)
    #         except ValueError:
    #             continue

    if not frames:
        print("No frames to save!")
        return None, None

    height, width = frames[0].shape[:2]
    fourcc = cv2.VideoWriter_fourcc(*'H264')
    video_filename = f"{stream_id}/{stream_id}_{fence_id}_{now_time_str}.mp4"
    video_path = f"{video_root}/{video_filename}"

    video_writer = cv2.VideoWriter(video_path, fourcc, fps, (width, height))
    for frame in frames:
        if frame.shape[1] != width or frame.shape[0] != height:
            frame = cv2.resize(frame, (width, height))
        video_writer.write(frame)
    video_writer.release()

    video_url = f"{base_url}/videos/{video_filename}"
    return video_url, video_path


def save_key_frames(stream_id, fence_id, frames, image_root='./images', base_url='127.0.0.1:5000'):
    """
    直接在 image_root 下保存第一帧和最后一帧图片，不创建子文件夹。
    文件名格式: {stream_id}_{fence_id}_{时间戳}_1.jpg / _2.jpg
    """
    now_time_str = datetime.now().strftime('%Y%m%d_%H%M%S')
    os.makedirs(f"{image_root}/{stream_id}", exist_ok=True)

    # 删除超过7天的旧文件
    # for file in os.listdir(image_root):
    #     file_path = os.path.join(image_root, file)
    #     if os.path.isfile(file_path):
    #         try:
    #             ts = file.split('_')[-2]  # 倒数第二个是时间戳
    #             file_time = datetime.strptime(ts, '%Y%m%d_%H%M%S')
    #             if now - file_time > timedelta(days=7):
    #                 os.remove(file_path)
    #         except ValueError:
    #             continue

    if not frames:
        print("No frames to save!")
        return None, None

    urls, paths = [], []

    # 按大小压缩
    # for i, frame in enumerate([frames[0], frames[-1]], start=1):
    #     # 压缩成30KB以内
    #     img_bytes = compress_to_30kb(frame, max_size_kb=30)
    #     filename = f"{stream_id}_{fence_id}_{now_time_str}_{i}.jpg"
    #     filepath = os.path.join(image_root, filename)
    #     with open(filepath, "wb") as f:
    #         f.write(img_bytes)
    #
    #     file_url = f"{base_url}/images/{filename}"
    #     urls.append(file_url)
    #     paths.append(filepath)
    #
    # return urls, paths

    # 按分辨率压缩
    for i, frame in enumerate([frames[0], frames[-1]], start=1):
        frame = resize_to_720p(frame)
        filename = f"{stream_id}_{fence_id}_{now_time_str}_{i}.jpg"
        filepath = f"{image_root}/{stream_id}/{filename}"
        success = cv2.imwrite(filepath, frame)
        if not success:
            print(f"{filepath}保存失败")
        file_url = f"{base_url}/{filepath}"
        urls.append(file_url)
        paths.append(filepath)

    return urls, paths


def compress_to_30kb(frame, max_size_kb=30):
    """
    将单帧图像压缩到指定大小以内
    :param frame: OpenCV BGR 图像
    :param max_size_kb: 压缩后的目标大小（KB）
    :return: 压缩后的字节数据
    """
    encode_param = [int(cv2.IMWRITE_JPEG_QUALITY), 95]  # 初始高质量
    success, encoded_img = cv2.imencode('.jpg', frame, encode_param)

    # 如果还超出大小，就逐步降低质量
    quality = 90
    while success and len(encoded_img) > max_size_kb * 1024 and quality > 5:
        encode_param = [int(cv2.IMWRITE_JPEG_QUALITY), quality]
        success, encoded_img = cv2.imencode('.jpg', frame, encode_param)
        quality -= 5

    return encoded_img.tobytes()


def resize_to_720p(frame):
    """将图像压缩到 720p 高度，保持宽高比"""
    h, w = frame.shape[:2]
    target_h = 720
    scale = target_h / h
    target_w = int(w * scale)
    resized = cv2.resize(frame, (target_w, target_h), interpolation=cv2.INTER_AREA)
    return resized


def chinese_to_pinyin(text):
    """把中文字符转为拼音，其它字符保持不变"""
    return ''.join(lazy_pinyin(text))
