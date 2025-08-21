import cv2
import numpy as np
from datetime import datetime, timedelta
import shutil
from matplotlib import pyplot as plt
from PIL import Image
import io
from pypinyin import lazy_pinyin
from docx import Document, oxml
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
import os
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.header import Header
from email.utils import formataddr
from email.mime.application import MIMEApplication
import base64
from pathlib import Path


def send_email_alert(message, contact_value, image_list=None, subject="视频报警通知"):
    """
    发送邮件报警（HTML正文，可附加多张图片）

    参数：
        message - 邮件正文（HTML格式字符串）
        contact_value - 收件人邮箱（字符串）
        image_list - 图片路径列表（可选）
    返回：
        True: 发送成功
        False: 发送失败
    """
    from_email = "576467179@qq.com"
    auth_code = "mirozaqvewotbdci"

    msg = MIMEMultipart()
    msg['From'] = formataddr(("报警系统", from_email))
    msg['To'] = contact_value
    msg['Subject'] = Header(subject, 'utf-8')
    msg.attach(MIMEText(message.replace("\n", "<br>"), 'html', 'utf-8'))

    if image_list:
        for img_path in image_list:
            if img_path and os.path.exists(img_path):
                with open(img_path, 'rb') as f:
                    part = MIMEApplication(f.read(), Name=os.path.basename(img_path))
                    part['Content-Disposition'] = f'attachment; filename="{os.path.basename(img_path)}"'
                    msg.attach(part)

    try:
        server = smtplib.SMTP_SSL("smtp.qq.com", 465)
        server.login(from_email, auth_code)
        server.sendmail(from_email, [contact_value], msg.as_string())
        server.quit()
        log("SUCCESS", f"邮件发送成功: {contact_value}")
        return True
    except Exception as e:
        log("FAIL", f"邮件发送失败: {contact_value}, 错误: {e}")
        return False


def save_report_to_docx(content: str, save_dir: str, filename: str, title: str = None, images: list = None, image_captions: list = None):
    """
    保存报告到 Word 文件（docx），带中文/英文字体设置、标题、正文、图片及描述。
    :param content: 正文内容
    :param save_dir: 保存目录
    :param filename: 文件名
    :param title: 标题文本
    :param images: 图片路径列表
    :param image_captions: 图片描述列表，长度应与 images 一致
    :return: 文件保存路径
    """
    try:
        os.makedirs(save_dir, exist_ok=True)
        doc = Document()

        # ===== 标题 =====
        if title:
            heading = doc.add_heading(level=1)
            run = heading.add_run(title)
            run.font.size = Pt(22)  # 小二
            run.bold = True
            run.font.name = 'Times New Roman'  # 英文
            run._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), '宋体')  # 中文
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            # 设置行距1.5倍
            heading.paragraph_format.line_spacing = 1.5
            heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中

        # ===== 正文 =====
        para = doc.add_paragraph()
        run = para.add_run(content)
        run.font.size = Pt(12)  # 小四
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), '宋体')
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        para.paragraph_format.line_spacing = 1.0  # 单倍行距
        para.paragraph_format.first_line_indent = Cm(0.74)  # 约2字符缩进
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # ===== 插入图片 =====
        if images:
            for idx, img_path in enumerate(images):
                if os.path.exists(img_path):
                    try:
                        pic = doc.add_picture(img_path, width=Cm(12))  # 宽度固定12cm
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        # 图片描述
                        if image_captions and idx < len(image_captions):
                            caption_para = doc.add_paragraph()
                            caption_run = caption_para.add_run(image_captions[idx])
                            caption_run.font.size = Pt(9)  # 小五
                            caption_run.font.name = 'Times New Roman'
                            caption_run._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), '宋体')
                            caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    except Exception as e:
                        log("FAIL", f"插入图片失败 {img_path}: {e}")

        # 保存文件
        filepath = os.path.join(save_dir, filename)
        doc.save(filepath)
        log("SUCCESS", f"Word 报告已保存: {filepath}")
        return filepath
    except Exception as e:
        log("FAIL", f"保存 Word 报告失败: {filename}, 错误: {e}")
        return None


def points_to_abs_points(frame, fences):
    height, width = frame.shape[:2]
    fence_points = []
    for fence in fences:
        points = fence.get('points', [])
        abs_points = [(int(p['x'] * width), int(p['y'] * height)) for p in points]
        fence_points.append(abs_points)
    return fence_points


class LogColors:
    INFO = "\033[94m"  # 蓝色
    WARNING = "\033[93m"  # 黄色
    FAIL = "\033[91m"  # 红色
    SUCCESS = "\033[92m"  # 绿色
    RESET = "\033[0m"  # 重置颜色


def log(level: str, message: str):
    """统一彩色日志打印，带白色时间戳"""
    color_map = {
        "INFO": LogColors.INFO,
        "WARNING": LogColors.WARNING,
        "FAIL": LogColors.FAIL,
        "SUCCESS": LogColors.SUCCESS
    }
    color = color_map.get(level, LogColors.INFO)
    # 获取当前时间戳，白色显示
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    WHITE = "\033[97m"
    RESET = LogColors.RESET
    print(f"{WHITE}{timestamp}{RESET} {color}[{level}]{RESET} {message}")


# 在报警分发时，给frame加上红色围栏标记
def draw_fence_on_frame(frame, fence_points):
    """
    在frame上绘制红色围栏（点和连线）
    fence_points: [(x1, y1), (x2, y2), ...] 电子围栏顶点像素坐标列表
    """
    if not fence_points or len(fence_points) < 3:
        return frame

    # 转成numpy数组方便绘制
    pts = np.array(fence_points, np.int32).reshape((-1, 1, 2))

    # 绘制多边形轮廓，红色，线宽2
    cv2.polylines(frame, [pts], isClosed=True, color=(0, 0, 255), thickness=2)

    # 绘制顶点为红色色小圆点
    for (x, y) in fence_points:
        cv2.circle(frame, (x, y), radius=4, color=(0, 0, 255), thickness=-1)

    return frame


def cv2_frame_to_base64(frame):
    # 转成 RGB
    rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
    pil_img = Image.fromarray(rgb_frame)
    buffered = io.BytesIO()
    pil_img.save(buffered, format="JPEG")
    img_str = base64.b64encode(buffered.getvalue()).decode()
    return img_str


def image_path_to_base64(image_path: str) -> str:
    if not os.path.exists(image_path):
        raise FileNotFoundError(f"文件不存在: {image_path}")
    with open(image_path, "rb") as f:
        img_bytes = f.read()
    return base64.b64encode(img_bytes).decode("utf-8")


# MD5哈希函数
def md5(str):
    import hashlib
    m = hashlib.md5()  # 创建MD5哈希对象
    m.update(str.encode("utf8"))  # 更新哈希内容
    return m.hexdigest()  # 返回16进制哈希值


def save_frames_as_video(stream_id, fence_id, frames, video_root='./videos', base_url='127.0.0.1:5000', fps=25):
    """
    直接在 video_root 下生成 MP4 文件，不创建子文件夹。
    文件名格式: {stream_id}_{fence_id}_{时间戳}.mp4
    """
    now_time_str = datetime.now().strftime('%Y%m%d_%H%M%S')
    now = datetime.now()
    Path(f"{video_root}/{stream_id}").mkdir(exist_ok=True)

    # 删除超过7天的旧文件
    # for file in os.listdir(video_root):
    #     file_path = os.path.join(video_root, file)
    #     if os.path.isfile(file_path):
    #         try:
    #             # 从文件名解析时间戳
    #             ts = file.split('_')[-1].split('.')[0]
    #             file_time = datetime.strptime(ts, '%Y%m%d_%H%M%S')
    #             if now - file_time > timedelta(days=7):
    #                 os.remove(file_path)
    #         except ValueError:
    #             continue

    if not frames:
        print("No frames to save!")
        return None, None

    height, width = frames[0].shape[:2]
    fourcc = cv2.VideoWriter_fourcc(*'avc1')
    video_filename = f"{stream_id}/{stream_id}_{fence_id}_{now_time_str}.mp4"
    video_path = os.path.join(video_root, video_filename)

    video_writer = cv2.VideoWriter(video_path, fourcc, fps, (width, height))
    for frame in frames:
        if frame.shape[1] != width or frame.shape[0] != height:
            frame = cv2.resize(frame, (width, height))
        video_writer.write(frame)
    video_writer.release()

    video_url = f"{base_url}/videos/{video_filename}"
    return video_url, video_path


def save_key_frames(stream_id, fence_id, frames, image_root='./images', base_url='127.0.0.1:5000'):
    """
    直接在 image_root 下保存第一帧和最后一帧图片，不创建子文件夹。
    文件名格式: {stream_id}_{fence_id}_{时间戳}_1.jpg / _2.jpg
    """
    now_time_str = datetime.now().strftime('%Y%m%d_%H%M%S')
    now = datetime.now()
    Path(f"{image_root}/{stream_id}").mkdir(exist_ok=True)

    # 删除超过7天的旧文件
    # for file in os.listdir(image_root):
    #     file_path = os.path.join(image_root, file)
    #     if os.path.isfile(file_path):
    #         try:
    #             ts = file.split('_')[-2]  # 倒数第二个是时间戳
    #             file_time = datetime.strptime(ts, '%Y%m%d_%H%M%S')
    #             if now - file_time > timedelta(days=7):
    #                 os.remove(file_path)
    #         except ValueError:
    #             continue

    if not frames:
        print("No frames to save!")
        return None, None

    urls, paths = [], []

    # 按大小压缩
    # for i, frame in enumerate([frames[0], frames[-1]], start=1):
    #     # 压缩成30KB以内
    #     img_bytes = compress_to_30kb(frame, max_size_kb=30)
    #     filename = f"{stream_id}_{fence_id}_{now_time_str}_{i}.jpg"
    #     filepath = os.path.join(image_root, filename)
    #     with open(filepath, "wb") as f:
    #         f.write(img_bytes)
    #
    #     file_url = f"{base_url}/images/{filename}"
    #     urls.append(file_url)
    #     paths.append(filepath)
    #
    # return urls, paths

    # 按分辨率压缩
    for i, frame in enumerate([frames[0], frames[-1]], start=1):
        frame = resize_to_720p(frame)
        filename = f"{stream_id}_{fence_id}_{now_time_str}_{i}.jpg"
        filepath = f"{image_root}/{stream_id}/{filename}"
        success = cv2.imwrite(filepath, frame)
        if not success:
            print(f"{filepath}保存失败")
        file_url = f"{base_url}/{filepath}"
        urls.append(file_url)
        paths.append(filepath)

    return urls, paths


def compress_to_30kb(frame, max_size_kb=30):
    """
    将单帧图像压缩到指定大小以内
    :param frame: OpenCV BGR 图像
    :param max_size_kb: 压缩后的目标大小（KB）
    :return: 压缩后的字节数据
    """
    encode_param = [int(cv2.IMWRITE_JPEG_QUALITY), 95]  # 初始高质量
    success, encoded_img = cv2.imencode('.jpg', frame, encode_param)

    # 如果还超出大小，就逐步降低质量
    quality = 90
    while success and len(encoded_img) > max_size_kb * 1024 and quality > 5:
        encode_param = [int(cv2.IMWRITE_JPEG_QUALITY), quality]
        success, encoded_img = cv2.imencode('.jpg', frame, encode_param)
        quality -= 5

    return encoded_img.tobytes()


def resize_to_720p(frame):
    """将图像压缩到 720p 高度，保持宽高比"""
    h, w = frame.shape[:2]
    target_h = 720
    scale = target_h / h
    target_w = int(w * scale)
    resized = cv2.resize(frame, (target_w, target_h), interpolation=cv2.INTER_AREA)
    return resized


def chinese_to_pinyin(text):
    """把中文字符转为拼音，其它字符保持不变"""
    return ''.join(lazy_pinyin(text))
